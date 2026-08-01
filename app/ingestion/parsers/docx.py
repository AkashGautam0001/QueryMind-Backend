import io
import docx
from app.ingestion.parsers.base import (
    DocumentParser, EmptyDocumentError, ParsedDocument, ParsedPage
)

class DocxParser(DocumentParser):
    def parse(self, content: bytes) -> ParsedDocument:
        try:
            doc = docx.Document(io.BytesIO(content))
        except Exception as exc:
            raise ValueError(f"Could not open file as DOCX: {exc}") from exc
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        if not full_text.strip():
            raise EmptyDocumentError("No extractable text found in this DOCX file.")

        return ParsedDocument(
            text=full_text,
            pages=[ParsedPage(page_number=None, text=full_text)],
            metadata={"paragraph_count" : len(paragraphs)}
        )
        